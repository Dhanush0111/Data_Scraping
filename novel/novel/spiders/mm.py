import math
import re
import scrapy
import docx

class Novelfull_net(scrapy.Spider):
    name = 'novelfull'
    allowed_domains = ['novelfulll.com', 'looknovel.com', 'lightnovelplus.com','watchnovel.com']  # Add all relevant domains
    start_urls = ['https://novelfulll.com/book/3443.html?page_num=3']
    bookname = 'I accidentally married a "CEO"'
    start = 101
    end = 200
    count = start
    first_page_found = False
    doc = docx.Document()
    
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS': 1,
        'RETRY_TIMES': 40,
        'DOWNLOAD_TIMEOUT': 180,  # Increase timeout
        'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
    }
    
    def parse(self, response):
        self.page_number = math.ceil(self.start / 50)
        urls = self.start_urls[0].split("?")[0]
        page_url = f'{urls}?page_num={self.page_number}'
        print("Page url is :", page_url)
        yield response.follow(url=page_url, callback=self.current_page)

    def current_page(self, response):
        links = response.css('.list-chapter li a::attr(href)').getall()
        print("Length of links is :", len(links))
        link_no = self.start % 50
        yield response.follow(url=links[link_no - 1], callback=self.parse_chapter)

    def parse_chapter(self, response):
        chapter_content = response.css('div.chapter-c p::text').getall()

        # Check if the first line starts with a number between 1 and 9
        if chapter_content:
            first_line = chapter_content[0].strip()
            if first_line and first_line[0].isdigit() and int(first_line[0]) in range(1, 10):
                # Remove the number at the beginning of the chapter
                chapter_content[0] = re.sub(r'^\d+\s*', '', first_line)
        
        # Remove extra blank lines
        chapter_content = [line.strip() for line in chapter_content if line.strip()]
                
        chapter_title = response.css('a.chapter-title span::text').get()
        print("Heading is :", chapter_title)
        self.doc.add_heading(chapter_title)
        
        for paragraph in chapter_content:
            self.doc.add_paragraph(paragraph)

        self.doc.add_page_break()

        next_btn = response.css('a.btn.btn-success#next_chap::attr(href)').get()
        next_btn = "https://novelfulll.com" + next_btn
        self.count += 1

        if self.count <= self.end:
            yield response.follow(next_btn, callback=self.parse_chapter)

    def closed(self, reason):
        self.doc.save(f'{self.bookname} {self.start} to {self.end}.docx')
