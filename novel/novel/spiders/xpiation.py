import scrapy
import docx
from bs4 import BeautifulSoup as bs

class Xpiaotianl(scrapy.Spider):
    name = 'xpiaotian'

    def __init__(self, **kwargs):
        self.name = 'The stunning poisonous doctor princess Row 645'   # Change this   
        self.book_url = 'https://www.xpiaotian.com/zh_hant/book/62014'   # Change this
        
        self.start_urls = [self.book_url]     
        self.start = 30 # Change this to the starting chapter number
        self.end = 32   # Change this to the ending chapter number
        self.count = self.start
        self.output_dict = {}
        self.doc = docx.Document()
        self.priority = 100000000
        self.website = ''
                
    custom_settings={
        # 'CONCURRENT_REQUESTS': 1
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS':1,
        'RETRY_TIMES': 10,
        # 'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        # 'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
        'DOWNLOAD_DELAY': 2, 
    }
    
    def parse(self, response):
        self.website = response.url.split('/book')[0]
        chapter_links = response.css('dd+ dt ~dd a::attr("href")').getall()
        first_chapter = chapter_links[self.start-1]
        yield response.follow(url=first_chapter, callback=self.forcontent, priority=self.priority)
        self.priority -= 1
    
    def forcontent(self, response):
        print(self.count)
        title = response.css('h1::text').get()
        soup = bs(response.text, "lxml")
        content = soup.select_one('#content')
        self.doc.add_heading(title)
        self.doc.add_heading("")
        for para in content:
            p = para.text.strip()
            if p == '':
                continue
            self.doc.add_paragraph(p)
            # self.doc.add_paragraph('')
        self.doc.add_page_break()
        next_url = soup.select_one('#A3').get('href')
        
        x = next_url.rsplit('/', 1)
        if '_' not in x[-1]:
            if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
                self.doc.save(f'{self.name}--{self.start if((self.count - 100)<self.start) else (self.count -100)} -{self.count}.docx')
                self.doc = docx.Document()
                self.start = self.count+1
            self.count += 1
        if self.count <= self.end:
            self.priority -= 1
            yield response.follow(url=next_url, callback=self.forcontent, priority=self.priority)
