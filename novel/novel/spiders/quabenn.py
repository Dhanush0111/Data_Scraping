import scrapy
import docx
import re

class zonghengSpider(scrapy.Spider):
    name = "quaben"
    allowed_domains = ["quanben-xiaoshuo.com"]
    start_urls = ["https://quanben-xiaoshuo.com/n/chengjia/1.html"]
    custom_settings = {
        'ROBOTSTXT_OBEY': False
    }
    def __init__(self):
        self.doc = docx.Document()
        self.start = 1
        self.end = 10
        self.name = "Mommy"  
        self.doc = docx.Document()
        self.priority = 100000000

    def parse(self, response):
        heading=response.css(".title::text").get()
        print("Heading is:", heading)
        content = response.css(".articlebody p::text").getall()
        

        # pattern = re.compile(r'\d+')
        # match = pattern.search(str(heading))
        # if match:
        #     # Get the matched integer value
        #     chapters_number = int(match.group())
        # #Add data to file
        self.doc.add_heading(heading, level=1)
        self.doc.add_heading("")
        for para in content:
            self.doc.add_paragraph(para)
        self.doc.add_page_break()

        try:
            # Code to save the document
            self.doc.save(f'{self.name}--{self.start} -{self.end}.docx')
        except Exception as e:
            self.logger.error(f"Error saving document: {e}")

        relative_url= response.css('.button::attr(href)').getall()[-1]
        next_chapter_url = "https://quanben-xiaoshuo.com/" +relative_url
        print("Next chapter URL is :", next_chapter_url)
        if self.start < self.end :
            yield response.follow(next_chapter_url, callback=self.parse)
    
    def closed(self,reason):
        print(reason)
        print("CLOSEDDDDDDDDDDDDDDDD")
        self.doc.save('trial.docx') 