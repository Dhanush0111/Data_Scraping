import scrapy
import re
import docx

class PowertoolContentEbtangSpider(scrapy.Spider):
    name = "Ebtang"
    # allowed_domains = ['www.ebtang.com', 'ebtang.com']

    # custom_settings = {
    #     'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
    #     'DOWNLOAD_HANDLERS' : {'https': 'scrapy.core.downloader.handlers.http10.HTTP10DownloadHandler'}
    # }
    custom_settings = {
    'ROBOTSTXT_OBEY': False,
    'CONCURRENT_REQUESTS': 1,
    'RETRY_TIMES': 40,
    'DOWNLOAD_TIMEOUT': 180,  # Increase timeout
    'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
    }
    def __init__(self, **kwargs):
        # self.start_urls = [kwargs['book_url']]
        # self.start = kwargs['start']
        # self.end = kwargs['end']
        self.start_urls = ["https://www.ebtang.com/book/3645"]
        self.start = 1
        self.end = 5
        # self.output_dict = kwargs['output_dict']
        self.counter = 1
        self.doc = docx.Document()
        self.total_chapters = None


    def parse(self, response):
        directory_url = response.url + '/directory'
        yield response.follow(directory_url, callback=self.parse_directory)

    def parse_directory(self, response):
        chapter_ids = response.css("b.chapter::attr('d-id')").getall()
        self.total_chapters = len(chapter_ids)

        first_chapter_link = f'{self.start_urls[0]}/{chapter_ids[0]}'
        yield response.follow(first_chapter_link, callback=self.parse_link_to_generate_randomStr, meta={"chapter_ids":chapter_ids})

    def parse_link_to_generate_randomStr(self, response):
        html = response.text
        chapter_ids = response.meta.get('chapter_ids')
        pattern = r'<input[^>]*id="randomString"[^>]*value="([^"]+)"'

        match = re.search(pattern, html)
        random_string = match.group(1)

        book_id = self.start_urls[0].split('/')[-1]
        first_chapter_id = chapter_ids[0]
        link = f"https://www.ebtang.com/book/{book_id}/{first_chapter_id}"
        yield response.follow(link, callback=self.parse_api, meta ={"first_chapter_id" : first_chapter_id, "book_id": book_id, "random_string": random_string})

    def parse_api(self, response):
        first_chapter_id = response.meta.get('first_chapter_id')
        book_id = response.meta.get('book_id')
        random_string = response.meta.get('random_string')

        next_chapter_id = first_chapter_id

        api_url = f'https://www.ebtang.com/book/readbook/{book_id}/{next_chapter_id}/more?randomString={random_string}'
        yield response.follow(api_url, callback=self.parse_content, meta={"book_id": book_id, "random_string": random_string})

    def parse_content(self, response):
        data = response.json()
        next_chapter_id = data['nextId']
        book_id = response.meta.get('book_id')
        random_string = response.meta.get("random_string")

        if self.counter >= self.start:
            try:
                title = data['bookChapter']['title']
                content = data['bookChapter']['content']
                print(content)
                self.doc.add_heading(title)
                self.doc.add_paragraph('')
                for p in content:
                    if p =='':
                        continue
                    p = p.strip()
                    self.doc.add_paragraph(p)
                self.doc.add_page_break()
            except Exception as e: # handles VIP chapters
                pass

        if self.counter <= self.end and self.counter <= self.total_chapters:
            self.counter += 1
            api_url = f'https://www.ebtang.com/book/readbook/{book_id}/{next_chapter_id}/more?randomString={random_string}'
            yield response.follow(api_url, callback=self.parse_content,meta={"book_id": book_id, "random_string": random_string})
        else:
            # self.output_dict['output'] = self.doc
            self.doc.save('EbtangFileTesting.docx')