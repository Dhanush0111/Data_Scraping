import math
import re
import scrapy
import docx

class Novelfull_net(scrapy.Spider):
    name = 'novelfulll123'
    allowed_domains = ['novelfulll.com']
    start_urls = ['https://novelfulll.com/book/3443.html']
    bookname = 'I accidentally married a "CEO"'
    start = 1
    end = 200
    count = start
    first_page_found = False
    doc = docx.Document()
    custom_settings={
        'CONCURRENT_REQUESTS':1,
        'ROBOTSTXT_OBEY': False,}
    
    def parse(self, response):
        self.page_number = math.ceil(self.start/50)
        if "?page_num=" in self.start_urls[0] :
            urls = self.start_urls[0].split("?")[0]
            page_url = f'{urls}?page_num={self.page_number}'
        else :
            urls = self.start_urls[0]
            page_url = f'{urls}?page_num={self.page_number}'
        yield response.follow(url=page_url,callback=self.parse_page, priority=self.priority)

    

    def current_page(self, response) :
        links = response.css('.list-chapter li a::attr(href)').getall()
        print("Length of links is :", len(links))
        link_no = self.start % 50
        yield response.follow(url=links[link_no - 1], callback=self.parse_chapter)

    def parse_chapter(self, response):
        chapter_title = response.css('a.chapter-title span::text').get()
        print("Heading is :", chapter_title)
        chapter_content1 = (response.css('div.chapter-c div::text').get()).strip()
        chapter_content = response.css('div.chapter-c p::text').getall()
        self.doc.add_heading(chapter_title)
        if chapter_content1 :
            self.doc.add_paragraph(chapter_content1)
        for paragraph in chapter_content:
            self.doc.add_paragraph(paragraph)

        self.doc.add_page_break()

        next_btn = response.css('a.btn.btn-success#next_chap::attr(href)').get()
        next_btn = "https://novelfulll.com" + next_btn
        self.count += 1

        if self.count <= self.end :
            yield response.follow(next_btn, callback=self.parse_chapter)

    def closed(self, reason):
        self.doc.save(f'{self.bookname} {self.start} to {self.end}.docx')