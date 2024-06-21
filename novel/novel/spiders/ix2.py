import scrapy
import time
from docx import Document
from urllib.parse import urlparse, parse_qs

class ixunshuSpider(scrapy.Spider) :
    name = 'ixunshu1'
    allowed_domains = ["ixunshu.net"]
    start_urls = ["https://ixunshu.net/xs/18288989/zj/193024974"]
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS':1,
        'RETRY_TIMES': 5,
        'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'DEFAULT_REQUEST_HEADERS': {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en',
        }
    }

    def __init__(self):
        self.start = 401
        self.end =500
        self.name = "After the princess was released from prison with her cub"  
        self.doc = Document()
        self.priority = 100000000

    def parse(self, response):
        if response.status == 403:
            self.logger.warning(f"Access to {response.url} forbidden, please check if the website allows scraping or requires additional headers/authentication.")
            return
        
        url_query = urlparse(response.url)
        query_params = parse_qs(url_query.query)
        
        # Check if the 'page' parameter is present in the URL
        if 'page' in query_params:
            content = response.css("#booktxt p::text").getall()
            for para in content:
                self.doc.add_paragraph(para)
        else:
            heading = response.css(".bookname::text").get()
            print("Heading is:", heading)
            self.doc.add_heading(heading, level=1)
            content = response.css("#booktxt p::text").getall()
            for para in content:
                self.doc.add_paragraph(para)
        
        next_chapter_url= response.css('.bottem2 a::attr(href)').getall()
        next_chapter_url = next_chapter_url[-1]
        print("Next chapter URL is :", next_chapter_url)
        
        time.sleep(20)
        if next_chapter_url != "/xs/18288989/zj/193025247" :
            self.priority -= 1000
            yield response.follow("https://ixunshu.net/" + next_chapter_url, callback=self.parse, priority=self.priority)
        else:
            self.doc.save(f'{self.name}-{self.start}-{self.end}.docx')

    def closed(self,reason):
        print(reason)
        print("CLOSEDDDDDDDDDDDDDDDD")
        self.doc.save(f'{self.name}-{self.start}-{self.end}.docx')
