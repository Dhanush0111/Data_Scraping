import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
import time
from docx import Document

# start_url = 'https://www.xiaomitemai.com/list/kan/595/595469/1'

start = 101
end = 299
book_name = "Crazy;  The princess blooms into a crazy flower"
firstChapter = (
    "https://booktoki341.com/novel/5072241?stx=%ED%99%A9%EB%85%80+%EB%AF%B8%EC%B9%9C+%EA%BD%83%EC%9C%BC%EB%A1%9C+%ED%94%BC%EC%96%B4%EB%82%98%EB%8B%A4&book=%EC%99%84%EA%B2%B0%EC%86%8C%EC%84%A4&spage=1"
)
start_chapter = start
current = start
oneDoc = 100   #no of chapters in a doc file

options = uc.ChromeOptions()
options.page_load_strategy = 'eager'
# options.add_argument('--headless')
driver = uc.Chrome(options=options)

driver.get(firstChapter)
time.sleep(10)
soup = bs(driver.page_source,"lxml")
time.sleep(0.5)
doc = Document()
while current<=end:
    if current>=start_chapter - 1:
        try:   
            print(current)
            heading = soup.select_one('.toon-title').text
            content = soup.select('#novel_content p')
            doc.add_heading(heading)
            for para in content:
                doc.add_paragraph(para.text)
            doc.add_page_break()
            nextLink = soup.select_one('#goNextBtn').get('href')
            
            driver.get(nextLink)
            time.sleep(1)
            soup = bs(driver.page_source,"lxml")
            current += 1

            if (current - start)%oneDoc == 0 and (current - start)!= -1 :
                doc.save(f"./output/{book_name} {start} to {current-1}.docx")      #  add './out/' before path for running in server
                start = current 
                doc = Document()
        except:
            break

doc.save(f"{book_name} {start} to {current}.docx")     #  add './out/' before path for running in server