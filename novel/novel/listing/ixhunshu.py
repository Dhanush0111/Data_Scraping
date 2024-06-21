import docx, time
import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
from selenium.webdriver.common.by import By


first_chpt_url = "https://ixunshu.net/xs/18288989/zj/193024794"   #start chapt link  #change here
book_name = "Born with no talent"     #change here
start = 1     #change here
end = 8      #change here
curr_page = start
doc = docx.Document()

options = uc.ChromeOptions()
# options.add_argument('--headless')
options.page_load_strategy = 'eager'
driver = uc.Chrome(options=options)

next_url=first_chpt_url
while(curr_page<=end):
    print('page: ',curr_page)
    driver.get(next_url)
    # time.sleep(2)
    soup=bs(driver.page_source,'lxml')
    heading=soup.select_one('.bookname').text
    doc.add_heading(heading.strip())
    paras=soup.select('#booktxt p')
    for c in paras:
        doc.add_paragraph(c.text.strip())
    doc.add_page_break()
    if((curr_page-start+1)%100==0 and curr_page!=end):
        doc.save(f'{book_name}-{start}-{curr_page}.docx')
        start=curr_page+1
        doc=docx.Document()
    elif curr_page==end:
        doc.save(f'{book_name}-{start}-{curr_page}.docx')
        break
    curr_page+=1
    next_url=driver.find_element(By.XPATH,"//a[div[@class='next-content']]").get_attribute('href')
    # print(next_url)
    # next_url=soup.select('.bottem2').get_attribute('href')
    next_url = soup.select('.bottem2')[0]['href']
    next_url="https://ixunshu.net/"+ next_url[-1]
driver.close()        



# next_chapter_url= response.css('.bottem2 a::attr(href)').getall()
#         next_chapter_url = next_chapter_url[-1]
#         print("Next chapter URL is :", next_chapter_url)