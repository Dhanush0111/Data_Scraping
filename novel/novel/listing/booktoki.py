import time
import re
from DrissionPage import ChromiumPage
from bs4 import BeautifulSoup
from docx import Document

# url = "https://www.vodtw.la/book/14077/chapter/12269931.html"


class Vodtw:
    def __init__(self, book_name, page_start, page_end):
        self.chapterNumber = page_start
        self.start_chapter = page_start
        self.oneDoc = 100
        self.end_chapter = page_end
        self.book_name = book_name
        self.doc = Document()
        self.page1 = ChromiumPage()

    def scrape(self, url):
        self.page1.get(url)
        time.sleep(5)
        print("URL->", url)
        try:
            soup = BeautifulSoup(self.page1.html, "lxml")
            # print(soup.prettify().encode("utf-8", errors="ignore"))
        except AttributeError as e:
            print("While making soup.", e)
            return  # Exit the function if there's an error

        chapter_title = soup.select_one(".toon-title").text
        chapter_title = re.sub(r'\(\d+/\d+\)', '', chapter_title)

        paragraphs = soup.select("#novel_content p")
        # if paragraphs:
        #     paragraphs_text = [p.get_text(strip=True) for p in paragraphs]
        #     paragraphs_text = '\n\n'.join(paragraphs_text)  # Join paragraphs with double newline
        # else:
        #     print("Paragraphs not found.No valid path found.")
            # return  # Exit the function if paragraphs not found
        # if paragraphs:
        #     paragraphs = paragraphs.get_text(strip=True)
        # else:
        #     print("Paragraphs not found.")
        #     # return  # Exit the function if paragraphs not found

        # try:
        #     self.doc.add_heading(chapter_title, level=1)
        #     self.doc.add_paragraph(paragraphs)
        #     # for text in paragraphs:
        #     #     self.doc.add_paragraph(text)
        #     self.doc.add_page_break()
        # except AttributeError as e:
        #     print("While Adding Heading and Paragraphs to doc file.", e)
        #     return  # Exit the function if there's an error
        self.doc.add_heading(chapter_title)
        # self.doc.add_heading('')
        for para in paragraphs:
            self.doc.add_paragraph(para.text.strip())
        self.doc.add_page_break()

        print("Crawled page number :", self.chapterNumber)

        # nextURL = soup.find("a", string="下一章")
        # if nextURL:
        #     nextURL = nextURL.get("href")
        #     nextPageURL = "https://www.vodtw.la/book/14077/chapter/" + nextURL
        #     print("Next url is :", nextPageURL)
        # else:
        #     print("Failed while fetching next URL.")
        #     return  # Exit the function if next URL not found
        # next_url_element = soup.find(
        #     "a",
        #     onclick=lambda value: value
        #     and "location.href=sn.replace('.idx','.html')" in value,
        # if next_url_element:
        #     next_url = next_url_element.get("href")
        #     print("Next URL:", next_url)
        # else:
        #     print("Failed to fetch the next URL.")
        #     return

        # working->
        # next_script_tag = soup.find(
        #     "script", text=lambda text: text and "var sn =" in text
        # )
        # if next_script_tag:

        #     script_text = next_script_tag.text

        #     start_index = script_text.find("var sn = '") + len("var sn = '")
        #     end_index = script_text.find("';", start_index)

        #     sn_value = script_text[start_index:end_index]
        #     next_chapter_link = (
        #         "https://www.vodtw.la/book/14077/chapter/"
        #         + sn_value.replace(".idx", ".html")
        #     )

        #     print("next_chapter_link", next_chapter_link)
        # till here

        # next_page_link = soup.select_one(
        #     "a.btn-comp[href*='javascript:move(\"next\")']"
        # )

        # # Check if the link exists
        # if next_page_link:
        #     # Extract the URL from the href attribute
        #     href = next_page_link.get("href")
        #     # Extract the URL parameters from the href attribute
        #     params = href.split("'")[1::2]
        #     # Construct the next page URL using the extracted parameters
        #     page_url = f"/?menu=novel&action=view&id={params[0]}"
        #     next_chapter_url = "https://www.munpia.com" + page_url
        next_chapter = soup.select_one('#goNextBtn').get('href')
        # print(next_chapter)
        # next_chapter_url="https://www.lightnovelworld.co" + next_chapter

        self.chapterNumber += 1

        if self.chapterNumber <= self.end_chapter:
            self.scrape(next_chapter)

    def save_document(self):
        file_name = (
            f"{self.book_name}_pages_{self.start_chapter}_to_{self.end_chapter}.docx"
        )
        try:
            self.doc.save(file_name)
            print(f"Document saved as '{file_name}'.")
        except Exception as e:
            print("Error saving document:", e)

    def close_driver(self):
        self.page1.quit()


if __name__ == "__main__":
    book_name = "Absolute swordsmanship"
    page_start = 201
    page_end = 299
    url = "https://booktoki341.com/novel/5072562?stx=%ED%99%A9%EB%85%80+%EB%AF%B8%EC%B9%9C+%EA%BD%83%EC%9C%BC%EB%A1%9C+%ED%94%BC%EC%96%B4%EB%82%98%EB%8B%A4&book=%EC%99%84%EA%B2%B0%EC%86%8C%EC%84%A4&spage=1"
    scraper = Vodtw(book_name, page_start, page_end)
    scraper.scrape(url)
    scraper.close_driver()
    scraper.save_document()