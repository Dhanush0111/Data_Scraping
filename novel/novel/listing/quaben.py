import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
import pandas as pd
import time
import docx
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
import math
import re
from selenium import webdriver

def read_number(str):
    pattern = r"第(.*?)章"
    lst = re.findall(pattern, str)
    try:
        return int(lst[0])
    except:
        return -1

doc = docx.Document()

options = uc.ChromeOptions()
options.add_argument("--disable-blink-features=AutomationControlled")
# options.add_argument('--headless')
options.add_argument(
    f"--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36"
)

# Specify the path to the downloaded ChromeDriver executable
chrome_driver_path = '/path/to/your/chromedriver'  # Update with your ChromeDriver path
driver = uc.Chrome(executable_path=chrome_driver_path, options=options)

book_chinese_name = "guifang"  # change here
start_chapter = 1  # change here
end_chapter = 4  # change here
site = f"https://www.quanben-xiaoshuo.com/n/{book_chinese_name}/xiaoshuo.html"

driver.get(site)
chpt_title_eles = driver.find_elements(By.CSS_SELECTOR, ".list a span")
chpt_titles = []
for ele in chpt_title_eles:
    chpt_titles.append(ele.text)
last_chapter = read_number(chpt_titles[-1])
chpt_link_eles = driver.find_elements(By.CSS_SELECTOR, ".list a")
chpt_links = []
for ele in chpt_link_eles:
    chpt_links.append(ele.get_attribute("href"))

i = 0
while i < len(chpt_links):
    if i + 1 < start_chapter:
        i += 1
        continue
    print(i + 1)
    driver.get(chpt_links[i])
    chpt_title = driver.find_element(By.CSS_SELECTOR, ".title").text
    doc.add_heading(chpt_title)
    chpt_content = driver.find_elements(By.CSS_SELECTOR, "p")
    for para_ele in chpt_content:
        para = para_ele.text.strip()
        if para:
            doc.add_paragraph(para)
    doc.add_page_break()
    if (i + 1) % 100 == 0:
        doc.save(f"quanben-{book_chinese_name}-{start_chapter}-{i+1}.docx")
        doc = docx.Document()
        start_chapter = i + 1 + 1
    i += 1

doc.save(f"quanben-{book_chinese_name}-{start_chapter}-{end_chapter}.docx")
