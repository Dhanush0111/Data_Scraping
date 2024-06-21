import scrapy
import docx
from bs4 import BeautifulSoup as bs

class PowertoolContentNovel543Spider(scrapy.Spider):
    name = 'powertool_content_novel543'
    
    def __init__(self, **kwargs):
        book_url = 'https://www.novel543.com/0823192343'

        if len(book_url.rsplit('/')[-1])==0:
            self.start_urls = [book_url + 'dir']
        else:
            self.start_urls = [book_url + '/dir']
        self.start = 50
        self.end = 60
        # self.output_dict = kwargs['output_dict']
        self.doc = docx.Document()
        self.priority = 1000000
        self.count = self.start -1
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response):
        chapter_links = response.css('.all a::attr("href")').getall()
        first_chapter = chapter_links[self.start-1]
        yield response.follow(url = first_chapter, callback=self.forcontent, priority = self.priority)
        self.priority -= 1
         
    def forcontent(self,response):
        title = response.css('h1::text').get()
        content = bs(response.text,"lxml").select('p')
        self.doc.add_heading(title)
        for para in content:
            self.doc.add_paragraph(para.text)
        self.doc.add_page_break()
    
        try:
            heads = title.rsplit('/',1)
            noOfPagesInChapter = int(heads[-1][0])
            current_page = int(heads[0][-1])
            if current_page==noOfPagesInChapter:
                self.count +=1
                # print(f'chapter {self.count} complete')
        except:
            noOfPagesInChapter = 1
            self.count +=1
            # print(f'chapter {self.count} complete (no parts)')

        if  self.count == self.end:
            # self.output_dict['output'] = self.doc
            self.doc.save(f'{self.start} -{self.count}.docx')

        if self.count<self.end:
            nextPageURL = response.css('.foot-nav a::attr("href")').getall()[2]
            # print(nextPageURL)
            yield response.follow(url = nextPageURL, callback=self.forcontent)