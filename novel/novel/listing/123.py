
mport scrapy
import time
from docx import Document

class ixunshuSpider(scrapy.Spider) :
    name = 'ixunshu'
    allowed_domains = ["ixunshu.net"]
    start_urls = ["https://ixunshu.net/xs/18288989/zj/193024794"]
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS':1,
        'RETRY_TIMES': 5,
        'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
    }

    def __init__(self):
        self.start = 289
        self.end =300
        self.name = "After the princess was released from prison with her cub"  
        self.doc = Document()
        self.priority = 100000000

    def parse(self, response):
        heading=response.css(".bookname::text").get()
        print("Heading is:", heading)
        self.doc.add_heading(heading, level=1)
        content = response.css("#booktxt p::text").getall()
        # print(content)
        for para in content :
            # print('Para is :', para)
            self.doc.add_paragraph(para)
        # self.doc.add_page_break()
            # Code to save the document
        
        next_chapter_url= response.css('.bottem2 a::attr(href)').getall()
        next_chapter_url = next_chapter_url[-1]
        print("Next chapter URL is :", next_chapter_url)
        
        time.sleep(20)
        if next_chapter_url != "/xs/18288989/zj/193024811" :
            self.priority -= 1000
            yield response.follow("https://ixunshu.net/" + next_chapter_url, callback=self.parse, priority=self.priority)
        else :
            self.doc.save(f'{self.name}-{self.start}-{self.end}.docx')

    def closed(self,reason):
        print(reason)